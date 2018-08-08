# coding:utf-8

'''
@author = super_fazai
@File    : __init__.py.py
@Time    : 2017/8/8 15:18
@connect : superonesfazai@gmail.com
'''

import docx

# 获取文档对象
file_path = '/Users/afa/Downloads/jj.docx'
file=docx.Document(file_path)
# print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段

# 输出每一段的内容
for para in file.paragraphs:
    print(para.text)

# 输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
#     print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)

for item in file.tables:
    print(item.__str__())

